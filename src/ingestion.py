import os
from pathlib import Path
import re
import concurrent.futures
import logging
from tqdm import tqdm

from langchain.docstore.document import Document  # Integration with LangChain Document objects

import PyPDF2
from bs4 import BeautifulSoup
from docx import Document as DocxDocument  # For DOCX file support

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    """Clean and standardize the input text."""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[str]:
    """Split text into chunks with a specified overlap."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - chunk_overlap
    return chunks

def load_text_file(file_path: str) -> str:
    """Load and return the content of a text file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def load_pdf_file(file_path: str) -> tuple[str, int]:
    """Extract text from a PDF file and return the text along with its page count."""
    text = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        page_count = len(reader.pages)
        for page in reader.pages:
            extracted = page.extract_text() or ""
            text.append(extracted)
    return " ".join(text), page_count

def load_html_file(file_path: str) -> str:
    """Extract visible text from an HTML file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')
        # Remove script and style tags
        for tag in soup(['script', 'style']):
            tag.extract()
        return soup.get_text(separator=' ')

def load_docx_file(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(file_path)
    return " ".join([para.text for para in doc.paragraphs])

def process_file(file_path: Path, chunk_size: int, chunk_overlap: int) -> list[Document]:
    """Process a single file: load, clean, chunk the text and wrap chunks in Document objects."""
    documents = []
    try:
        raw_text = ""
        metadata = {"source": str(file_path)}
        if file_path.suffix.lower() == '.txt':
            raw_text = load_text_file(str(file_path))
        elif file_path.suffix.lower() == '.pdf':
            raw_text, page_count = load_pdf_file(str(file_path))
            metadata["page_count"] = page_count
        elif file_path.suffix.lower() in ['.html', '.htm']:
            raw_text = load_html_file(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            raw_text = load_docx_file(str(file_path))
        else:
            # Skip unknown file types
            return []

        cleaned = clean_text(raw_text)
        chunks = chunk_text(cleaned, chunk_size, chunk_overlap)
        for chunk in chunks:
            documents.append(Document(page_content=chunk, metadata=metadata))
    except Exception as e:
        logger.error(f"Error processing {file_path}: {e}")
    return documents

def ingest_documents(data_dir: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[Document]:
    """
    Ingest documents from a given directory.
    
    This function scans the directory for supported file types, processes them in parallel,
    and returns a list of LangChain Document objects with extracted text and metadata.
    """
    documents = []
    data_path = Path(data_dir)
    file_paths = [fp for fp in data_path.rglob('*') if fp.is_file()]

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {
            executor.submit(process_file, file_path, chunk_size, chunk_overlap): file_path
            for file_path in file_paths
        }
        for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="Processing files"):
            try:
                docs = future.result()
                documents.extend(docs)
            except Exception as e:
                logger.error(f"Error processing file {futures[future]}: {e}")

    return documents

# Example usage:
if __name__ == "__main__":
    data_directory = "./data"  # Configure your data directory here
    docs = ingest_documents(data_directory)
    print(f"Ingested {len(docs)} document chunks.")
